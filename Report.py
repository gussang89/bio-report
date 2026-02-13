import streamlit as st
import urllib.request
import urllib.parse
import json
import re
import google.generativeai as genai
from datetime import datetime, timedelta
import io
from docx import Document
import xml.etree.ElementTree as ET
from email.utils import parsedate_to_datetime

# --- 1. 구글 제미나이 설정 ---
def configure_gemini():
    if "GOOGLE_API_KEY" in st.secrets:
        genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
        return True
    return False

# --- 2. 검색 함수들 ---

# [1] 해외 논문 (Europe PMC)
def get_epmc_papers(keywords, months):
    query_parts = [f'({k.strip()})' for k in keywords if k.strip()]
    if not query_parts: return []
    keyword_query = " OR ".join(query_parts)
    end_date = datetime.now()
    start_date = end_date - timedelta(days=months*30)
    date_query = f'FIRST_PDATE:[{start_date.strftime("%Y-%m-%d")} TO {end_date.strftime("%Y-%m-%d")}]'
    full_query = f"({keyword_query}) AND ({date_query})"
    encoded_query = urllib.parse.quote(full_query)
    base_url = f"https://www.ebi.ac.uk/europepmc/webservices/rest/search?query={encoded_query}&format=json&resultType=core&pageSize=30"
    
    try:
        req = urllib.request.Request(base_url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as response:
            data = json.loads(response.read().decode('utf-8'))
        filtered = []
        for p in data.get('resultList', {}).get('result', []):
            title = p.get('title', '')
            abstract = re.sub('<[^<]+>', '', p.get('abstractText', ''))
            doi = p.get('doi')
            link = f"https://doi.org/{doi}" if doi else ""
            if title and abstract:
                filtered.append({"title": title, "abstract": abstract, "url": link, "date": p.get('firstPublicationDate', '')})
        return filtered
    except: return []

# [2] 핵심 수정: 국내 뉴스 (네이버 뉴스 RSS 적용 - API 키 불필요)
def get_domestic_news(keywords, months):
    news_list = []
    headers = {'User-Agent': 'Mozilla/5.0'}
    cutoff_date = datetime.now() - timedelta(days=months*30)
    
    for k in keywords:
        clean_k = k.strip()
        if not clean_k: continue
        
        encoded_query = urllib.parse.quote(clean_k)
        # 네이버 뉴스 검색 RSS URL
        url = f"https://newssearch.naver.com/search.naver?where=rss&query={encoded_query}"
        
        try:
            req = urllib.request.Request(url, headers=headers)
            with urllib.request.urlopen(req) as response:
                root = ET.fromstring(response.read())
            
            for item in root.findall('./channel/item'):
                title = item.find('title').text
                link = item.find('link').text
                pubDate = item.find('pubDate').text
                description = item.find('description').text
                
                # 날짜 파싱 및 기간 필터링
                try: 
                    dt = parsedate_to_datetime(pubDate)
                    # 설정한 기간 이전의 뉴스는 버림
                    if dt.replace(tzinfo=None) < cutoff_date.replace(tzinfo=None):
                        continue
                    date_str = dt.strftime("%Y-%m-%d")
                except: 
                    date_str = pubDate
                
                # 제목과 초록의 불필요한 HTML 태그 깔끔하게 제거
                clean_title = re.sub('<[^<]+>', '', title)
                clean_abstract = re.sub('<[^<]+>', '', description) if description else "상세 내용은 링크 참고"
                
                news_list.append({"title": clean_title, "abstract": clean_abstract[:300], "url": link, "date": date_str})
        except Exception as e:
            st.warning(f"'{clean_k}' 네이버 검색 중 오류: {e}")
            continue
            
    unique_news = {n['url']: n for n in news_list}.values()
    return sorted(unique_news, key=lambda x: x['date'], reverse=True)

# [3] 핵심 수정: 해외 뉴스 (구글 뉴스 버그 우회)
def get_overseas_news(keywords, months):
    news_list = []
    headers = {'User-Agent': 'Mozilla/5.0'}
    cutoff_date = datetime.now() - timedelta(days=months*30)
    
    for k in keywords:
        clean_k = k.strip()
        if not clean_k: continue
        
        encoded_query = urllib.parse.quote(clean_k)
        # 구글 서버 버그를 일으키는 when 옵션을 빼고, 파이썬에서 날짜를 걸러냅니다.
        url = f"https://news.google.com/rss/search?q={encoded_query}&hl=en-US&gl=US&ceid=US:en"
        
        try:
            req = urllib.request.Request(url, headers=headers)
            with urllib.request.urlopen(req) as response:
                root = ET.fromstring(response.read())
            
            for item in root.findall('./channel/item'):
                title = item.find('title').text
                link = item.find('link').text
                pubDate = item.find('pubDate').text
                
                try: 
                    dt = parsedate_to_datetime(pubDate)
                    if dt.replace(tzinfo=None) < cutoff_date.replace(tzinfo=None):
                        continue
                    date_str = dt.strftime("%Y-%m-%d")
                except: 
                    date_str = pubDate
                    
                news_list.append({"title": title, "abstract": "상세 내용은 원문 참조", "url": link, "date": date_str})
        except Exception as e:
            continue
            
    unique_news = {n['url']: n for n in news_list}.values()
    return sorted(unique_news, key=lambda x: x['date'], reverse=True)

# --- 3. AI 리포트 생성 ---
def generate_ai_report(items, keywords, section_type):
    if not items: return "분석할 데이터가 없습니다."
    
    data_text = ""
    for i, item in enumerate(items[:30]):
        data_text += f"[{i+1}] 제목: {item['title']} (일자: {item['date']})\n초록: {item['abstract'][:200]}\n\n"

    if section_type == "Paper":
        role = "글로벌 바이오 공정 연구원"
        focus = "최신 공정 기술, 수율 개선, 촉매 동향 분석"
    elif section_type == "Domestic_News":
        role = "한국 바이오 에너지 시장 애널리스트"
        focus = "국내 정책 변화, 정유/바이오 기업의 동향, 규제 흐름"
    else:
        role = "글로벌 바이오 에너지 시장 애널리스트"
        focus = "해외 선진국의 상용화 동향, 주요 규제, 글로벌 기업 투자 동향"

    prompt = f"""
    당신은 {role}입니다. 키워드: {', '.join(keywords)}
    
    아래 데이터를 바탕으로 **'심층 보고서'**를 작성하세요.
    
    [핵심 준수 사항: 정확성 및 근거 표기]
    1. 모든 서술은 반드시 제공된 데이터를 근거로 해야 하며, 문장이나 단락 끝에 반드시 출처 주석(예: [1], [3])을 달아주세요.
    2. 데이터만으로 명확히 알 수 없어 논리적으로 추론하거나 애매한 부분에 대해서는, 반드시 **"※ 추론: 본 내용은 명시된 데이터가 부족하여 문맥을 바탕으로 추론된 것으로 정확성에 한계가 있을 수 있습니다."**라고 서술하세요. 없는 내용을 절대 지어내지 마세요.

    [작성 포인트]
    1. 📊 **핵심 트렌드 요약**: {focus}
    2. 💡 **세부 분석 및 인사이트**: 주요 이슈 및 현업 적용/대응 시사점
    3. 📌 **주요 원문 리뷰**: 핵심 데이터 번호 기재하여 리뷰

    [수집된 데이터]
    {data_text}
    """
    
    try:
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        preferred_order = ['models/gemini-1.5-pro', 'models/gemini-1.5-flash', 'models/gemini-pro']
        sorted_models = [m for m in preferred_order if m in available_models] + [m for m in available_models if m not in preferred_order]

        for model_name in sorted_models:
            try:
                model = genai.GenerativeModel(model_name)
                response = model.generate_content(prompt)
                return response.text
            except: continue
        return "AI 분석 실패."
    except: return "모델 설정 오류."

# --- 4. Word 생성 ---
def create_word_doc(report_text, keywords, title):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"검색 키워드: {', '.join(keywords)}")
    doc.add_paragraph("-" * 50)
    for line in report_text.split('\n'):
        if line.startswith('###'): doc.add_heading(line.replace('###', '').strip(), level=3)
        elif line.startswith('##'): doc.add_heading(line.replace('##', '').strip(), level=2)
        elif line.startswith('#'): doc.add_heading(line.replace('#', '').strip(), level=1)
        elif line.startswith('**') and line.endswith('**'): 
            p = doc.add_paragraph()
            p.add_run(line.replace('**', '')).bold = True
        else: doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 5. 메인 UI ---
st.set_page_config(page_title="Bio-Energy Tracker", layout="wide")
st.title("🔬 바이오 에너지 통합 트래커 (네이버 뉴스 탑재)")
st.caption("각 탭에서 원하는 주제의 검색 버튼을 눌러 개별적으로 리포트를 생성하세요.")

if not configure_gemini():
    st.error("❌ Google API Key 설정 필요")

with st.sidebar:
    st.header("🔍 카테고리별 검색어 설정")
    
    st.subheader("1. 해외 논문 (영어)")
    paper_keywords = st.text_area("공정/기술 키워드", value="Biodiesel production\nTransesterification catalyst", height=100)
    
    st.subheader("2. 국내 뉴스 (한글)")
    domestic_keywords = st.text_area("국내 시장/정책 키워드", value="바이오디젤\n지속가능항공유\n에쓰오일 바이오\nHD현대오일뱅크 바이오", height=100)
    
    st.subheader("3. 해외 뉴스 (영어)")
    overseas_keywords = st.text_area("해외 시장/정책 키워드", value="Sustainable Aviation Fuel\nHVO market\nNeste biofuel", height=100)
    
    st.divider()
    months = st.slider("검색 기간 (개월)", 1, 24, 6)

tab_paper, tab_domestic, tab_overseas = st.tabs(["🌍 논문 분석 (해외 기술)", "🇰🇷 국내 뉴스 분석 (네이버)", "🌎 해외 뉴스 분석"])

with tab_paper:
    st.markdown("### 🌍 해외 바이오 공정 기술 탐색")
    if st.button("해외 논문 검색 및 분석 🚀", key="btn_run_paper"):
        k_paper = [k.strip() for k in paper_keywords.split('\n') if k.strip()]
        if not k_paper: st.warning("검색어를 사이드바에 입력해주세요.")
        else:
            with st.spinner("해외 논문을 수집 및 분석 중입니다..."):
                papers = get_epmc_papers(k_paper, months)
                if not papers: st.warning("검색된 해외 논문이 없습니다.")
                else:
                    st.success(f"성공! {len(papers)}건의 논문을 바탕으로 리포트를 작성했습니다.")
                    report_paper = generate_ai_report(papers, k_paper, "Paper")
                    docx_paper = create_word_doc(report_paper, k_paper, "🌍 바이오 논문/기술 분석 리포트")
                    st.download_button("📥 논문 리포트 다운로드", docx_paper, "Paper_Report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="btn_dl_paper")
                    st.divider()
                    st.markdown(report_paper)
                    with st.expander("📝 수집된 논문 원문 리스트 보기"):
                        for i, p in enumerate(papers):
                            st.write(f"**[{i+1}] {p['title']}** ({p['date']})  [링크]({p['url']})")

with tab_domestic:
    st.markdown("### 🇰🇷 국내 바이오 시장 및 정책 탐색")
    if st.button("국내 뉴스 검색 및 분석 🚀", key="btn_run_domestic"):
        k_domestic = [k.strip() for k in domestic_keywords.split('\n') if k.strip()]
        if not k_domestic: st.warning("검색어를 사이드바에 입력해주세요.")
        else:
            with st.spinner("네이버에서 국내 뉴스를 수집 및 분석 중입니다..."):
                d_news = get_domestic_news(k_domestic, months)
                if not d_news: st.warning("검색된 국내 뉴스가 없습니다. 검색어를 바꿔보세요.")
                else:
                    st.success(f"성공! {len(d_news)}건의 뉴스를 바탕으로 리포트를 작성했습니다.")
                    report_domestic = generate_ai_report(d_news, k_domestic, "Domestic_News")
                    docx_domestic = create_word_doc(report_domestic, k_domestic, "🇰🇷 국내 바이오 시장/정책 리포트")
                    st.download_button("📥 국내 뉴스 리포트 다운로드", docx_domestic, "Domestic_News_Report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="btn_dl_domestic")
                    st.divider()
                    st.markdown(report_domestic)
                    with st.expander("📝 수집된 국내 뉴스 원문 리스트 보기"):
                        for i, n in enumerate(d_news):
                            st.write(f"**[{i+1}] {n['title']}** ({n['date']})  [링크]({n['url']})")

with tab_overseas:
    st.markdown("### 🌎 해외 바이오 시장 및 정책 탐색")
    if st.button("해외 뉴스 검색 및 분석 🚀", key="btn_run_overseas"):
        k_overseas = [k.strip() for k in overseas_keywords.split('\n') if k.strip()]
        if not k_overseas: st.warning("검색어를 사이드바에 입력해주세요.")
        else:
            with st.spinner("해외 뉴스를 수집 및 분석 중입니다..."):
                o_news = get_overseas_news(k_overseas, months)
                if not o_news: st.warning("검색된 해외 뉴스가 없습니다. 검색어를 바꿔보세요.")
                else:
                    st.success(f"성공! {len(o_news)}건의 뉴스를 바탕으로 리포트를 작성했습니다.")
                    report_overseas = generate_ai_report(o_news, k_overseas, "Overseas_News")
                    docx_overseas = create_word_doc(report_overseas, k_overseas, "🌎 해외 바이오 시장/정책 리포트")
                    st.download_button("📥 해외 뉴스 리포트 다운로드", docx_overseas, "Overseas_News_Report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="btn_dl_overseas")
                    st.divider()
                    st.markdown(report_overseas)
                    with st.expander("📝 수집된 해외 뉴스 원문 리스트 보기"):
                        for i, n in enumerate(o_news):
                            st.write(f"**[{i+1}] {n['title']}** ({n['date']})  [링크]({n['url']})")
